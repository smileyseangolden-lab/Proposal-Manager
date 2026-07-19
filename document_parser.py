"""Parse uploaded RFP/RFQ documents into plain text for the agent."""

from pathlib import Path

import docx
import PyPDF2

from config.settings import VERTICALS


def parse_document(file_path: str, ocr: bool = False) -> str:
    """Parse a document file and return its text content.

    Supports PDF, DOCX, plain text, and Excel formats. (Excel matters because
    RFQs frequently arrive as bid forms / pricing schedules in .xlsx — the
    upload UI advertises it, so silently skipping it would drop scope.)

    When ``ocr=True`` and OCR is available, a PDF whose embedded text is sparse
    (a scan) gets an OCR pass, and image files are OCR'd directly. OCR is slow,
    so callers pass ``ocr=True`` only from background jobs — never the request
    path.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = _parse_pdf(path)
        if ocr and len(text.strip()) < 100:
            try:
                import ocr as _ocr
                ocr_text = _ocr.pdf_text(str(path))
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
            except Exception:
                pass
        return text
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path)
    elif suffix in (".txt", ".md", ".csv"):
        return path.read_text(encoding="utf-8", errors="replace")
    elif suffix in (".xlsx", ".xls"):
        from rate_sheet_parser import parse_rate_sheet
        return parse_rate_sheet(str(path)).get("raw_text", "")
    elif ocr and suffix in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif"):
        import ocr as _ocr
        return _ocr.image_text(str(path))
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def _parse_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    text_parts = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _parse_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def detect_vertical(text: str) -> str:
    """Keyword-scoring vertical detection — the FALLBACK classifier.

    proposal_agent.classify_vertical() is the primary path (a fast LLM call);
    it delegates here when no API key is configured or the call fails, so this
    must stay dependency-free and never raise. Returns the vertical key
    (e.g. 'data_center'), or 'general' if no strong match is found.
    """
    text_lower = text.lower()

    # Data Center / Mission Critical signals
    dc_keywords = [
        "data center", "data centre", "datacenter", "mission critical",
        "colocation", "colo ", "hyperscale", "bms", "epms",
        "building management system", "electrical power monitoring",
        "ups ", "uninterruptible power", "cooling system",
        "crah", "data hall", "megawatt", "pue ",
        "raised floor", "hot aisle", "cold aisle", "containment",
        "generator", "switchgear", "pdu ", "power distribution",
        "ats ", "automatic transfer", "leak detection",
    ]
    dc_score = sum(text_lower.count(kw) for kw in dc_keywords)

    # Life Science / Pharma signals
    ls_keywords = [
        "pharmaceutical", "pharma", "life science", "biopharmaceutical",
        "gmp", "good manufacturing practice", "fda ", "21 cfr",
        "cleanroom", "clean room", "iso class", "aseptic",
        "validation", "iq/oq/pq", "installation qualification",
        "operational qualification", "performance qualification",
        "eu annex", "batch record", "environmental monitoring",
        "biotech", "bioreactor", "fill finish", "lyophiliz",
        "sterile", "cip ", "sip ", "wfi ", "purified water",
    ]
    ls_score = sum(text_lower.count(kw) for kw in ls_keywords)

    # Food & Beverage / CPG signals
    fb_keywords = [
        "food and beverage", "food & beverage", "food processing",
        "consumer packaged good", "cpg ", "beverage",
        "haccp", "fsma", "food safety", "sqf ",
        "gfsi", "sanitary", "hygienic design", "washdown",
        "usda", "packaging line", "bottling", "canning",
        "pasteuriz", "steriliz", "cold storage", "refrigerat",
        "food grade", "allergen", "traceability",
    ]
    fb_score = sum(text_lower.count(kw) for kw in fb_keywords)

    scores = {
        "data_center": dc_score,
        "life_science": ls_score,
        "food_beverage": fb_score,
    }

    best = max(scores, key=scores.get)
    # Require a minimum threshold to avoid false positives
    if scores[best] >= 3:
        return best
    return "general"


def load_vertical_resources(vertical_key: str) -> dict:
    """Load workflow, templates, and reference proposals for a specific vertical.

    Returns a dict with keys: 'workflow', 'templates', 'reference_proposals'.
    """
    vertical = VERTICALS.get(vertical_key)
    if not vertical:
        vertical = VERTICALS["general"]

    vertical_dir = vertical["dir"]

    # Load workflow
    workflow_path = vertical_dir / "workflow.md"
    workflow = ""
    if workflow_path.exists():
        workflow = workflow_path.read_text(encoding="utf-8")

    # Load templates
    templates: dict[str, str] = {}
    templates_dir = vertical_dir / "templates"
    if templates_dir.exists():
        for file_path in sorted(templates_dir.iterdir()):
            if file_path.suffix.lower() in (".md", ".txt"):
                templates[file_path.name] = file_path.read_text(encoding="utf-8")

    # Load reference proposals
    reference_proposals: list[str] = []
    ref_dir = vertical_dir / "reference_proposals"
    if ref_dir.exists():
        for file_path in sorted(ref_dir.iterdir()):
            if file_path.suffix.lower() in (".md", ".txt", ".pdf", ".docx"):
                try:
                    reference_proposals.append(parse_document(str(file_path)))
                except Exception:
                    continue

    return {
        "workflow": workflow,
        "templates": templates,
        "reference_proposals": reference_proposals,
    }


def load_reference_documents(reference_dir: Path) -> dict[str, list[str]]:
    """Load all reference documents organized by category.

    Returns a dict with keys: 'sample_rfps', 'sample_rfqs', 'past_proposals'
    each mapping to a list of document text contents.
    """
    categories = {
        "sample_rfps": reference_dir / "sample_rfps",
        "sample_rfqs": reference_dir / "sample_rfqs",
        "past_proposals": reference_dir / "past_proposals",
    }
    result: dict[str, list[str]] = {}
    for category, cat_dir in categories.items():
        docs = []
        if cat_dir.exists():
            for file_path in sorted(cat_dir.iterdir()):
                if file_path.suffix.lower() in (".md", ".txt", ".pdf", ".docx"):
                    try:
                        docs.append(parse_document(str(file_path)))
                    except Exception:
                        continue
        result[category] = docs
    return result


def load_templates(templates_dir: Path) -> dict[str, str]:
    """Load all proposal templates/boilerplate.

    Returns a dict mapping template filename to its text content.
    """
    templates: dict[str, str] = {}
    if templates_dir.exists():
        for file_path in sorted(templates_dir.iterdir()):
            if file_path.suffix.lower() in (".md", ".txt"):
                templates[file_path.name] = file_path.read_text(encoding="utf-8")
    return templates
