"""Export generated proposals to DOCX format."""

import difflib
import re
from datetime import datetime, timezone
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def markdown_to_docx(
    markdown_text: str,
    output_path: str,
    logo_path: str | None = None,
    logo_placement: str = "top_left",
    logo_on_cover: bool = False,
    company_name: str = "",
    font_name: str = "Calibri",
) -> str:
    """Convert a Markdown proposal to a formatted DOCX file.

    Args:
        markdown_text: The proposal in Markdown format.
        output_path: File path for the output DOCX.
        logo_path: Optional path to a company logo image to embed.
        logo_placement: "top_left" or "center" — where to put the logo.
        logo_on_cover: If True, add a dedicated cover page before the body.
        company_name: Company name displayed on the cover page (if enabled).
        font_name: Default body font to use.

    Returns:
        The output file path.
    """
    doc = docx.Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name or "Calibri"
    font.size = Pt(11)

    # Embed company logo
    logo_exists = bool(logo_path) and Path(logo_path).exists()
    if logo_exists:
        if logo_on_cover:
            _add_cover_page(doc, logo_path, logo_placement, company_name)
        else:
            _add_inline_logo(doc, logo_path, logo_placement)

    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Flush table if we leave a table block
        if in_table and not stripped.startswith("|"):
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading = doc.add_heading(stripped[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("").add_run().add_break()

        # Table rows
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Skip separator rows like |---|---|
            if all(re.match(r"^[-:]+$", c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")

        # Numbered lists
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")

        # Bold text as standalone line
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True

        # Empty line
        elif not stripped:
            pass  # Skip empty lines (spacing handled by styles)

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

        i += 1

    # Flush any remaining table
    if table_rows:
        _add_table(doc, table_rows)

    doc.save(output_path)
    return output_path


def _add_table(doc, rows: list[list[str]]):
    """Add a table to the document."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                table.rows[r_idx].cells[c_idx].text = cell_text

    doc.add_paragraph()  # spacing after table


def _add_formatted_text(paragraph, text: str):
    """Add text with basic inline formatting (bold, italic) to a paragraph."""
    # Split on bold markers
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _logo_width_for_placement(placement: str) -> Inches:
    """Return an appropriate rendered width based on placement."""
    if placement == "center":
        return Inches(3.0)
    return Inches(1.75)  # top_left — letterhead size


def _add_inline_logo(doc, logo_path: str, placement: str):
    """Add the logo at the very top of the document (before any content)."""
    p = doc.add_paragraph()
    if placement == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run()
    try:
        run.add_picture(str(logo_path), width=_logo_width_for_placement(placement))
    except Exception:
        # If the image can't be embedded, silently fall through — don't break export
        pass


def _add_cover_page(doc, logo_path: str, placement: str, company_name: str = ""):
    """Add a dedicated cover page with the logo, followed by a page break."""
    # Spacer paragraphs to vertically push logo off the top edge
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    if placement == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run()
    try:
        # Larger logo on dedicated cover pages
        width = Inches(4.0) if placement == "center" else Inches(2.5)
        run.add_picture(str(logo_path), width=width)
    except Exception:
        pass

    if company_name:
        p2 = doc.add_paragraph()
        if placement == "center":
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cn_run = p2.add_run(company_name)
        cn_run.font.size = Pt(16)
        cn_run.bold = True

    # Add a page break so the proposal body starts on a new page
    doc.add_page_break()


def markdown_to_redline_docx(original_md: str, revised_md: str, output_path: str,
                              author: str = "Reviewer") -> str:
    """Generate a DOCX with tracked changes (redlines) showing differences
    between original and revised markdown content.

    Uses visual redlining: deleted text is shown in red strikethrough,
    inserted text is shown in blue underline. This approach works reliably
    across all Word versions and doesn't require accepting/rejecting changes.

    Args:
        original_md: The original markdown text (e.g., AI version).
        revised_md: The revised markdown text (e.g., human-edited version).
        output_path: File path for the output DOCX.
        author: Name of the person who made the changes.

    Returns:
        The output file path.
    """
    doc = docx.Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading("Proposal — Tracked Changes", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    run = info.add_run(f"Changes by: {author} | Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Legend
    legend = doc.add_paragraph()
    del_run = legend.add_run("Red strikethrough = deleted  ")
    del_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    del_run.font.strike = True
    del_run.font.size = Pt(10)
    ins_run = legend.add_run("Blue underline = inserted")
    ins_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    ins_run.font.underline = True
    ins_run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

    # Split into lines and diff
    orig_lines = original_md.splitlines()
    rev_lines = revised_md.splitlines()
    differ = difflib.SequenceMatcher(None, orig_lines, rev_lines)

    for tag, i1, i2, j1, j2 in differ.get_opcodes():
        if tag == "equal":
            for line in orig_lines[i1:i2]:
                _add_redline_line(doc, line, "equal")
        elif tag == "delete":
            for line in orig_lines[i1:i2]:
                _add_redline_line(doc, line, "delete")
        elif tag == "insert":
            for line in rev_lines[j1:j2]:
                _add_redline_line(doc, line, "insert")
        elif tag == "replace":
            # For replacements, show word-level diffs within each line pair
            for idx in range(max(i2 - i1, j2 - j1)):
                old_line = orig_lines[i1 + idx] if (i1 + idx) < i2 else ""
                new_line = rev_lines[j1 + idx] if (j1 + idx) < j2 else ""

                if old_line and new_line:
                    _add_word_diff_line(doc, old_line, new_line)
                elif old_line:
                    _add_redline_line(doc, old_line, "delete")
                elif new_line:
                    _add_redline_line(doc, new_line, "insert")

    doc.save(output_path)
    return output_path


def _add_redline_line(doc, line: str, change_type: str):
    """Add a full line with redline formatting."""
    stripped = line.strip()
    if not stripped:
        return

    # Handle headings
    heading_level = 0
    if stripped.startswith("#### "):
        heading_level = 4
        stripped = stripped[5:]
    elif stripped.startswith("### "):
        heading_level = 3
        stripped = stripped[4:]
    elif stripped.startswith("## "):
        heading_level = 2
        stripped = stripped[3:]
    elif stripped.startswith("# "):
        heading_level = 1
        stripped = stripped[2:]

    if heading_level:
        p = doc.add_heading("", level=heading_level)
    else:
        p = doc.add_paragraph()

    run = p.add_run(stripped)
    if change_type == "delete":
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        run.font.strike = True
    elif change_type == "insert":
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.underline = True


def _add_word_diff_line(doc, old_line: str, new_line: str):
    """Add a paragraph with word-level diff highlighting."""
    stripped_old = old_line.strip()
    stripped_new = new_line.strip()

    # Check for heading
    heading_level = 0
    for prefix, level in [("#### ", 4), ("### ", 3), ("## ", 2), ("# ", 1)]:
        if stripped_new.startswith(prefix):
            heading_level = level
            stripped_old = stripped_old[len(prefix):] if stripped_old.startswith(prefix) else stripped_old
            stripped_new = stripped_new[len(prefix):]
            break

    if heading_level:
        p = doc.add_heading("", level=heading_level)
    else:
        p = doc.add_paragraph()

    old_words = stripped_old.split()
    new_words = stripped_new.split()
    sm = difflib.SequenceMatcher(None, old_words, new_words)

    for tag, a1, a2, b1, b2 in sm.get_opcodes():
        if tag == "equal":
            p.add_run(" ".join(old_words[a1:a2]) + " ")
        elif tag == "delete":
            run = p.add_run(" ".join(old_words[a1:a2]) + " ")
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            run.font.strike = True
        elif tag == "insert":
            run = p.add_run(" ".join(new_words[b1:b2]) + " ")
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
            run.font.underline = True
        elif tag == "replace":
            run = p.add_run(" ".join(old_words[a1:a2]) + " ")
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            run.font.strike = True
            run = p.add_run(" ".join(new_words[b1:b2]) + " ")
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
            run.font.underline = True


def markdown_to_rfi_docx(
    items: list,
    project_name: str,
    client_name: str,
    company_name: str,
    author: str,
    output_path: str,
) -> str:
    """Generate a formatted RFI/Clarification letter as DOCX.

    Args:
        items: List of ClarificationItem model instances (customer-facing).
        project_name: Name of the project.
        client_name: Name of the client/customer.
        company_name: Sender company name.
        author: Name of the person generating the letter.
        output_path: File path for the output DOCX.

    Returns:
        The output file path.
    """
    doc = docx.Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Header
    title = doc.add_heading("Request for Information / Clarification", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Letter metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Date: {datetime.now(timezone.utc).strftime('%B %d, %Y')}\n")
    meta.add_run(f"From: {company_name}\n")
    meta.add_run(f"To: {client_name or '[Customer Name]'}\n")
    meta.add_run(f"Re: {project_name}\n")
    meta.add_run(f"Prepared by: {author}\n")

    doc.add_paragraph()

    # Introduction
    intro = doc.add_paragraph()
    intro.add_run(
        f"Dear {client_name or '[Customer Name]'},\n\n"
        f"During our review of the RFP/RFQ documentation for the above-referenced project, "
        f"we have identified the following items requiring clarification. We respectfully "
        f"request your response to each item below to ensure our proposal accurately addresses "
        f"your requirements.\n\n"
        f"Please reference the RFI number in your response for tracking purposes."
    )

    doc.add_paragraph()

    # RFI Items
    doc.add_heading("Clarification Items", level=2)

    # Group by category
    categories = {}
    for item in items:
        cat = (item.category or "general").title()
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(item)

    for cat_name, cat_items in categories.items():
        doc.add_heading(cat_name, level=3)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"

        header_cells = table.rows[0].cells
        for i, hdr in enumerate(["RFI #", "Priority", "Question", "Response"]):
            header_cells[i].text = hdr
            for paragraph in header_cells[i].paragraphs:
                for r in paragraph.runs:
                    r.bold = True

        for item in cat_items:
            row = table.add_row()
            row.cells[0].text = item.rfi_reference_id or "-"
            row.cells[1].text = (item.priority or "medium").title()
            row.cells[2].text = item.question or ""
            row.cells[3].text = ""

        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(
        "We appreciate your prompt attention to these items. Please do not hesitate "
        "to contact us if you require any additional information.\n\n"
    )
    closing = footer.add_run(f"Sincerely,\n{author}\n{company_name}")
    closing.bold = True

    doc.save(output_path)
    return output_path
